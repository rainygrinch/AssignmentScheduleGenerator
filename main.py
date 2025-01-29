### VERSION RELEASE NOTES ###

### VERSION 1.1 ###
# THE FOLLOWING UPDATES HAVE BEEN MADE...
# Asks the user to specify the input file to be used, instead of specifying a file path
# This version also asks the user where they would like to save the output word files to
# Seeks a .csv file instead of a .xlsx file
# Ignores the first line of the csv export from Bullhorn
# Placeholders changed to use BH defaults (by updating the placeholders on the "template.docx" file

### VERSION 1.2 ###

# Added count for total attempted doc generation
# Added count for successful doc generation
# List unsuccessful doc generation by Placement ID
# Added welcome, disclaimer and instructions

### VERSION 1.3 ###

# Added process timer
# Removed Static Template File Directory (allow user to select themselves)

### VERSION 1.4 ###

# Completely refactored for GUI Interface
# Added progress bar and doc count in real time
# Fixed issue where successful doc gen count is always 1 less than total
# Corrected issue where failed IDs did not display

### VERSION 1.5 ###

# Added program information to main window, equivalent to title screen in
# version 1.3 and earlier
# BugFix: "Generate Docs" button does not display as pushed below window height
# Simplified title info
# Increased window size to allow for all content
# Added NOTICE PERIOD to TEMPLATE and CSV INPUT FILE

import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.ttk import Progressbar
from docx import Document
import time
import pandas as pd

class AssignmentScheduleGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Assignment Schedule Generator")
        self.root.geometry("500x500")  # Adjusted height to fit everything

        # Initialize variables
        self.input_file_path = None
        self.template_file_path = None
        self.output_folder = None
        self.attempted_count = 0
        self.successful_count = 0
        self.total_documents = 0
        self.failed_ids = []
        self.status_label = None

        # Create GUI components (buttons, labels, etc.)
        self.create_widgets()

    def create_widgets(self):
        # Display title information in the main window
        self.display_title_info()

        # File selection and status label
        self.input_button = tk.Button(self.root, text="Select Input File", command=self.select_input_file)
        self.input_button.grid(row=1, column=0, padx=10, pady=10, sticky="w")

        self.input_status_label = tk.Label(self.root, text="Not Selected", font=("Helvetica", 10))
        self.input_status_label.grid(row=1, column=1, padx=10, pady=10, sticky="w")

        self.template_button = tk.Button(self.root, text="Select Template File", command=self.select_template_file)
        self.template_button.grid(row=2, column=0, padx=10, pady=10, sticky="w")

        self.template_status_label = tk.Label(self.root, text="Not Selected", font=("Helvetica", 10))
        self.template_status_label.grid(row=2, column=1, padx=10, pady=10, sticky="w")

        self.output_button = tk.Button(self.root, text="Select Output Folder", command=self.select_output_folder)
        self.output_button.grid(row=3, column=0, padx=10, pady=10, sticky="w")

        self.output_status_label = tk.Label(self.root, text="Not Selected", font=("Helvetica", 10))
        self.output_status_label.grid(row=3, column=1, padx=10, pady=10, sticky="w")

        self.status_label = tk.Label(self.root, text="Status: Waiting for file selections...", font=("Helvetica", 12))
        self.status_label.grid(row=4, column=0, columnspan=2, pady=20)

        self.generate_button = tk.Button(self.root, text="Generate Documents", command=self.generate_assignment_schedules, state=tk.DISABLED)
        self.generate_button.grid(row=5, column=0, columnspan=2, pady=20)

    def display_title_info(self):
        # Title info as plain text
        title_text = """
        Assignment Schedule Generator v1.5
             Created by: Peter Grint
         For use by: IDPP Consulting Ltd Only
                www.idpp.com

  This program helps generate assignment schedules
   by replacing placeholders in a template file.

       DISCLAIMER: This is a prototype.
    No liability accepted for malfunctions."""

        # Display the title info in the window
        title_label = tk.Label(self.root, text=title_text, font=("Courier", 10), justify="left", padx=10, pady=10)
        title_label.grid(row=0, column=0, columnspan=2, pady=10)

    def select_input_file(self):
        self.input_file_path = filedialog.askopenfilename(title="Select the CSV Input file", filetypes=[("CSV Files", "*.csv")])
        if self.input_file_path:
            self.input_status_label.config(text="Ready")
            self.check_all_files_selected()

    def select_template_file(self):
        self.template_file_path = filedialog.askopenfilename(title="Select Template File", filetypes=[("Word Files", "*.docx")])
        if self.template_file_path:
            self.template_status_label.config(text="Ready")
            self.check_all_files_selected()

    def select_output_folder(self):
        self.output_folder = filedialog.askdirectory(title="Select the folder to save the output documents")
        if self.output_folder:
            self.output_status_label.config(text="Ready")
            self.check_all_files_selected()

    def check_all_files_selected(self):
        if self.input_file_path and self.template_file_path and self.output_folder:
            self.status_label.config(text="Status: Ready for document generation")
            self.generate_button.config(state=tk.NORMAL)

    def generate_assignment_schedules(self):
        if not self.input_file_path or not self.template_file_path or not self.output_folder:
            messagebox.showerror("Error", "Please select all the necessary files and folders.")
            return

        df = pd.read_csv(self.input_file_path, skiprows=[1])

        # Create progress window
        progress_window = tk.Toplevel(self.root)
        progress_window.title("Progress")
        progress_window.geometry("400x200")

        # Progress label and progress bar
        progress_label = tk.Label(progress_window, text="Generating documents...", font=("Helvetica", 12))
        progress_label.pack(pady=10)

        progress_bar = Progressbar(progress_window, length=300, mode="determinate")
        progress_bar.pack(pady=10)

        self.total_documents = len(df)
        progress_bar["maximum"] = self.total_documents

        document_count_label = tk.Label(progress_window, text="Generated: 0 / " + str(self.total_documents), font=("Helvetica", 10))
        document_count_label.pack(pady=10)

        # Track the start time of document generation
        start_time = time.time()  # Start time for elapsed time calculation

        # Initialise the list to track failed document IDs
        failed_ids = []

        # Close button function to be used after process completion
        def close_program():
            self.root.quit()

        # Proceed with all documents
        for index, row in df.iterrows():  # Loop through all rows in the CSV
            self.attempted_count += 1

            doc = Document(self.template_file_path)

            # Replace placeholders in paragraphs
            for para in doc.paragraphs:
                for column_name in df.columns:
                    placeholder = f"{{{{{column_name}}}}}"
                    replacement_value = str(row[column_name])

                    if placeholder in para.text:
                        for run in para.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, replacement_value)

            # Replace placeholders in tables
            for table in doc.tables:
                for row_cells in table.rows:
                    for cell in row_cells.cells:
                        for column_name in df.columns:
                            placeholder = f"{{{{{column_name}}}}}"
                            replacement_value = str(row[column_name])

                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, replacement_value)

            output_file = f"{self.output_folder}/Assignment_Schedule_{row['ID']}.docx"
            try:
                doc.save(output_file)
                self.successful_count += 1
            except Exception as e:
                failed_ids.append(row['ID'])  # Add failed document ID to the list

            # Update progress bar and count label
            progress_bar["value"] = self.attempted_count
            document_count_label.config(text=f"Generated: {self.attempted_count} / {self.total_documents}")
            progress_window.update_idletasks()

        # Finalisation after the process
        end_time = time.time()
        elapsed_time = end_time - start_time  # Calculate elapsed time

        # Update status and show summary
        self.status_label.config(text=f"Attempted: {self.attempted_count} | Success: {self.successful_count} | Time: {elapsed_time:.2f} seconds")

        # Show failed document IDs if any
        if failed_ids:
            failed_ids_str = ", ".join(map(str, failed_ids))
            messagebox.showinfo("Document Generation Summary",
                                f"Document generation completed.\n\n"
                                f"Failed IDs: {failed_ids_str}")
        else:
            messagebox.showinfo("Success", "Document generation completed with no failures.")

        # Add Close button only when the generation is finished
        close_button = tk.Button(progress_window, text="Close Program", command=close_program)
        close_button.pack(pady=10)

        progress_window.mainloop()

# Main part to run the GUI
if __name__ == "__main__":
    root = tk.Tk()
    app = AssignmentScheduleGeneratorApp(root)
    root.mainloop()
