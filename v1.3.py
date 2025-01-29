### VERSION RELEASE NOTES ###

### VERSION 1.1 ###
# THE FOLLOWING UPDATES HAVE BEEN MADE...
# Asks the user to specify the input file to be used, instead of specifying a file path
# This version also asks the user where they would like to save the output word files to
# Seeks a .csv file instead of a .xlsx file
# Ignores the first line of the csv export from Bullhorn
# Placeholders changed to use BH defaults (by updating the placeholders on the "template.docx" file

### VERSION 1.2 ###

# Added count for total attempted doc generation
# Added count for successful doc generation
# List unsuccessful doc generation by Placement ID
# Added welcome, disclaimer and instructions

### VERSION 1.3 ###

# Added process timer
# Removed Static Template File Directory (allow user to select themselves)

import time
import pandas as pd
from docx import Document
from tkinter import filedialog, Tk

def display_title_screen():
    width = 60  # Width of the title screen
    stars = '*' * width  # Full line of stars for the border

    title = f"""
{stars}
{' ' * ((width - len("Assignment Schedule Generator v1.2")) // 2)}Assignment Schedule Generator
{' ' * ((width - len("Created by: Peter Grint")) // 2)}Created by: Peter Grint
{' ' * ((width - len("For use by: IDPP Consulting Only")) // 2)}For use by: IDPP Consulting Only
{stars}
{' ' * ((width - len("This program helps generate assignment schedules")) // 2)}This program helps generate assignment schedules
{' ' * ((width - len("by replacing placeholders in a template file.")) // 2)}by replacing placeholders in a template file.
{stars}
{' ' * ((width - len("DISCLAIMER: This is a prototype.")) // 2)}DISCLAIMER: This is a prototype.
{' ' * ((width - len("No liability accepted for malfunctions.")) // 2)}No liability accepted for malfunctions.
{stars}
    """

    # Print the title screen
    print(title)


# Function to generate assignment schedules with success/failure reporting
def generate_assignment_schedules_with_report():
    # Ask for input CSV file
    root = Tk()
    root.withdraw()  # Don't need the root window
    print()
    print("***** STEP 1 *****")
    print("Please select the CSV file holding the placement data")
    print("A Windows Explorer box has opened, it might be on another screen")
    input_file_path = filedialog.askopenfilename(title="Select the CSV Input file", filetypes=[("CSV Files", "*.csv")])

    if not input_file_path:
        print("No CSV file selected. Exiting.")
        return

    # Read CSV file into DataFrame, skipping second row (index 1)
    df = pd.read_csv(input_file_path, skiprows=[1])

    # Ask for the template file location
    print("**THANK YOU** - Input File Received - **THANK YOU**")
    print()
    print("***** STEP 2 *****")
    print("Please select the template file (Word document) to use.")
    print("A Windows Explorer box has opened, it might be on another screen")
    template_file_path = filedialog.askopenfilename(title="Select Template File", filetypes=[("Word Files", "*.docx")])

    if not template_file_path:
        print("No template file selected. Exiting.")
        return

    # Ask for output folder
    print("**THANK YOU** - Template File Identified - **THANK YOU**")
    print()
    print("***** STEP 3 *****")
    print("Please select the folder you would like to save the Assignment")
    print("Schedules to - A Windows Explorer box has opened,")
    print("it might be on another screen")
    output_folder = filedialog.askdirectory(title="Select the folder to save the output documents")

    if not output_folder:
        print("No output folder selected. Exiting.")
        return

    # Start Timer
    start_time = time.time()

    # Variables to track success and failure counts
    failed_ids = []  # List to hold failed IDs
    attempted_count = 0  # Number of attempts
    successful_count = 0  # Number of successful generations

    # Loop through the DataFrame and generate a Word document for each row
    for index, row in df.iterrows():
        attempted_count += 1
        print(f"Processing row {index}")

        # Open the template document
        doc = Document(template_file_path)

        # Loop through paragraphs in the document and replace placeholders
        for para in doc.paragraphs:
            for column_name in df.columns:
                placeholder = f"{{{{{column_name}}}}}"  # Create the placeholder with curly braces
                replacement_value = str(row[column_name])  # Get value from the row (convert to string)

                if placeholder in para.text:
                    # Replace placeholder in the paragraph's text
                    inline = para.runs
                    for run in inline:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement_value)

        # Loop through tables in the document and replace placeholders
        for table in doc.tables:
            for row_cells in table.rows:  # row_cells are not iterable, we need to use row.cells
                for cell in row_cells.cells:  # Now iterating over cell in row
                    for column_name in df.columns:
                        placeholder = f"{{{{{column_name}}}}}"  # Create the placeholder with curly braces
                        replacement_value = str(row[column_name])  # Get value from the row (convert to string)

                        if placeholder in cell.text:
                            # Replace placeholder in the table cell text
                            cell.text = cell.text.replace(placeholder, replacement_value)

        # Save the modified document
        output_file = f"{output_folder}/Assignment_Schedule_{row['ID']}.docx"
        try:
            doc.save(output_file)
            print(f"Document saved: {output_file}")
            successful_count += 1
        except Exception as e:
            print(f"Failed to generate document for ID {row['ID']}. Error: {e}")
            failed_ids.append(row['ID'])

    # Stop Timer
    end_time = time.time()

    # Calculate elapsed time
    elapsed_time = end_time - start_time

    # Output results after processing all rows
    print(f"\nAttempted to generate {attempted_count} assignment schedules.")
    print(f"Successfully generated {successful_count} schedules.")
    if failed_ids:
        print(f"The following IDs failed to generate schedules: {', '.join(map(str, failed_ids))}")
    else:
        print("All assignment schedules were generated successfully.")

    return elapsed_time # make elapsed time available outside of function




# Display Title Screen
display_title_screen()

# Run function and get elapsed time
elapsed_time = generate_assignment_schedules_with_report()

# Display elapsed time if function completed successfully

if elapsed_time is not None:
    print(f"\nProgram execution complete in {elapsed_time:.2f} seconds.")

input("\nProgram execution completed. Press Enter to exit...")
