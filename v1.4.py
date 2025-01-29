### VERSION RELEASE NOTES ###

### VERSION 1.1 ###
# THE FOLLOWING UPDATES HAVE BEEN MADE...
# Asks the user to specify the input file to be used, instead of specifying a file path
# This version also asks the user where they would like to save the output word files to
# Seeks a .csv file instead of a .xlsx file
# Ignores the first line of the csv export from Bullhorn
# Placeholders changed to use BH defaults (by updating the placeholders on the "template.docx" file

### VERSION 1.2 ###

# Added count for total attempted doc generation
# Added count for successful doc generation
# List unsuccessful doc generation by Placement ID
# Added welcome, disclaimer and instructions

### VERSION 1.3 ###

# Added process timer
# Removed Static Template File Directory (allow user to select themselves)

### VERSION 1.4 ###

# Completely refactored for GUI Interface
# Added progress bar and doc count in real time
# Fixed issue where successful doc gen count is always 1 less than total
# Corrected issue where failed IDs did not display

### VERSION 1.5 ###

# Added program information (equivalent to title screen from v1.3 and earlier)
#


import os
import time
import pandas as pd
from docx import Document
from tkinter import filedialog, Tk, messagebox, Button, Label, StringVar, Frame, Toplevel
from tkinter.ttk import Progressbar

class DocumentGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Generator")
        self.root.geometry("600x400")

        # Initialize variables
        self.input_file_path = ""
        self.template_file_path = ""
        self.output_folder = ""
        self.attempted_count = 0
        self.successful_count = 0
        self.failed_ids = []
        self.total_documents = 0

        # GUI components
        self.create_widgets()

    def create_widgets(self):
        frame = Frame(self.root)
        frame.pack(pady=20)

        # Instructions Label
        self.instructions_label = Label(frame, text="Select Input File, Template, and Output Folder", font=("Helvetica", 12))
        self.instructions_label.grid(row=0, column=0, columnspan=2, pady=10)

        # Input File Button
        self.input_button = Button(frame, text="Select CSV Input File", command=self.select_input_file)
        self.input_button.grid(row=1, column=0, padx=10)

        self.input_file_label = Label(frame, text="No file selected", font=("Helvetica", 10))
        self.input_file_label.grid(row=1, column=1)

        # Template File Button
        self.template_button = Button(frame, text="Select Template File", command=self.select_template_file)
        self.template_button.grid(row=2, column=0, padx=10)

        self.template_file_label = Label(frame, text="No file selected", font=("Helvetica", 10))
        self.template_file_label.grid(row=2, column=1)

        # Output Folder Button
        self.output_button = Button(frame, text="Select Output Folder", command=self.select_output_folder)
        self.output_button.grid(row=3, column=0, padx=10)

        self.output_folder_label = Label(frame, text="No folder selected", font=("Helvetica", 10))
        self.output_folder_label.grid(row=3, column=1)

        # Status Label
        self.status_label = Label(frame, text="Status: Waiting for input files...", font=("Helvetica", 10))
        self.status_label.grid(row=4, column=0, columnspan=2, pady=20)

        # Start Button
        self.start_button = Button(frame, text="Start Document Generation", command=self.generate_assignment_schedules)
        self.start_button.grid(row=5, column=0, columnspan=2)

    def select_input_file(self):
        self.input_file_path = filedialog.askopenfilename(title="Select the CSV Input file", filetypes=[("CSV Files", "*.csv")])
        if self.input_file_path:
            self.input_file_label.config(text=f"{os.path.basename(self.input_file_path)}")
            self.update_status()

    def select_template_file(self):
        self.template_file_path = filedialog.askopenfilename(title="Select Template File", filetypes=[("Word Files", "*.docx")])
        if self.template_file_path:
            self.template_file_label.config(text=f"{os.path.basename(self.template_file_path)}")
            self.update_status()

    def select_output_folder(self):
        self.output_folder = filedialog.askdirectory(title="Select the folder to save the output documents")
        if self.output_folder:
            self.output_folder_label.config(text=f"{os.path.basename(self.output_folder)}")
            self.update_status()

    def update_status(self):
        if self.input_file_path and self.template_file_path and self.output_folder:
            self.status_label.config(text="Status: Ready for document generation")
        else:
            self.status_label.config(text="Status: Waiting for input files...")

    def generate_assignment_schedules(self):
        if not self.input_file_path or not self.template_file_path or not self.output_folder:
            messagebox.showerror("Error", "Please select all the necessary files and folders.")
            return

        df = pd.read_csv(self.input_file_path, skiprows=[1])

        # Create progress window
        progress_window = Toplevel(self.root)
        progress_window.title("Progress")
        progress_window.geometry("400x150")

        # Progress label and progress bar
        progress_label = Label(progress_window, text="Generating documents...", font=("Helvetica", 12))
        progress_label.pack(pady=10)

        progress_bar = Progressbar(progress_window, length=300, mode="determinate")
        progress_bar.pack(pady=10)

        self.total_documents = len(df)
        progress_bar["maximum"] = self.total_documents

        document_count_label = Label(progress_window, text="Generated: 0 / " + str(self.total_documents),
                                     font=("Helvetica", 10))
        document_count_label.pack(pady=10)

        # Track the start time of document generation
        start_time = time.time()  # Start time for elapsed time calculation

        # Initialize the list to track failed document IDs
        failed_ids = []

        # Close button function to be used after process completion
        def close_program():
            self.root.quit()

        # Proceed with all documents
        for index, row in df.iterrows():  # Loop through all rows in the CSV
            self.attempted_count += 1

            doc = Document(self.template_file_path)

            # Replace placeholders in paragraphs
            for para in doc.paragraphs:
                for column_name in df.columns:
                    placeholder = f"{{{{{column_name}}}}}"
                    replacement_value = str(row[column_name])

                    if placeholder in para.text:
                        for run in para.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, replacement_value)

            # Replace placeholders in tables
            for table in doc.tables:
                for row_cells in table.rows:
                    for cell in row_cells.cells:
                        for column_name in df.columns:
                            placeholder = f"{{{{{column_name}}}}}"
                            replacement_value = str(row[column_name])

                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, replacement_value)

            output_file = f"{self.output_folder}/Assignment_Schedule_{row['ID']}.docx"
            try:
                doc.save(output_file)
                self.successful_count += 1
            except Exception as e:
                failed_ids.append(row['ID'])  # Add failed document ID to the list

            # Update progress bar and count label
            progress_bar["value"] = self.attempted_count
            document_count_label.config(text=f"Generated: {self.attempted_count} / {self.total_documents}")
            progress_window.update_idletasks()

        # Finalization after the process
        end_time = time.time()
        elapsed_time = end_time - start_time  # Calculate elapsed time

        # Update status and show summary
        self.status_label.config(
            text=f"Attempted: {self.attempted_count} | Success: {self.successful_count} | Time: {elapsed_time:.2f} seconds")

        # Show failed document IDs if any
        if failed_ids:
            failed_ids_str = ", ".join(map(str, failed_ids))
            messagebox.showinfo("Document Generation Summary",
                                f"Document generation completed.\n\n"
                                f"Failed IDs: {failed_ids_str}")
        else:
            messagebox.showinfo("Success", "Document generation completed with no failures.")

        # Add Close button only when the generation is finished
        close_button = Button(progress_window, text="Close Program", command=close_program)
        close_button.pack(pady=10)

        progress_window.mainloop()


if __name__ == "__main__":
    root = Tk()
    app = DocumentGeneratorApp(root)
    root.mainloop()
