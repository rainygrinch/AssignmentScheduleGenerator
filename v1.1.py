#### VERSION 1.1 ####
# THE FOLLOWING UPDATES HAVE BEEN MADE...
# Asks the user to specify the input file to be used, instead of specifying a file path
# This version also asks the user where they would like to save the output word files to
# Seeks a .csv file instead of a .xlsx file
# Ignores the first line of the csv export from Bullhorn
# Placeholders changed to use BH defaults (by updating the placeholders on the "template.docx" file

import pandas as pd
from docx import Document
from tkinter import filedialog, Tk

# Static template path
template_file = r"C:\Users\petergrint.DPP_DOMAIN\PycharmProjects\ContractGenerator\template.docx"


# Function to generate assignment schedules
def generate_assignment_schedules():
    # Ask for input CSV file
    root = Tk()
    root.withdraw()  # Don't need the root window
    input_file_path = filedialog.askopenfilename(title="Select the CSV Input file", filetypes=[("CSV Files", "*.csv")])

    if not input_file_path:
        print("No CSV file selected. Exiting.")
        return

    # Read CSV file into DataFrame, skipping second row (index 1)
    df = pd.read_csv(input_file_path, skiprows=[1])

    # Display DataFrame preview for debugging
   # print("DataFrame preview:")
   # print(df.head())

    # Debug: Print column names to ensure they are read correctly
   # print("CSV Column names:", df.columns)

    # Ask for output folder
    output_folder = filedialog.askdirectory(title="Select the folder to save the output documents")

    if not output_folder:
        print("No output folder selected. Exiting.")
        return

    # Loop through the DataFrame and generate a Word document for each row
    for index, row in df.iterrows():
        print(f"Processing row {index}")

        # Open the template document
        doc = Document(template_file)

        # Loop through paragraphs and tables in the document
        for para in doc.paragraphs:
            for column_name in df.columns:
                placeholder = f"{{{{{column_name}}}}}"  # Create the placeholder with curly braces
                replacement_value = str(row[column_name])  # Get value from the row (convert to string)

                # Debug: Print the placeholder being checked and replaced
                # print(f"Checking placeholder: {placeholder} with value: {replacement_value}")

                if placeholder in para.text:
                    # Replace placeholder in the paragraph's text
                    inline = para.runs
                    for run in inline:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement_value)

        # Loop through tables in the document
        for table in doc.tables:
            for row_cells in table.rows:  # row_cells are not iterable, we need to use row.cells
                for cell in row_cells.cells:  # Now iterating over cell in row
                    for column_name in df.columns:
                        placeholder = f"{{{{{column_name}}}}}"  # Create the placeholder with curly braces
                        replacement_value = str(row[column_name])  # Get value from the row (convert to string)

                        # Debug: Print the placeholder being checked and replaced
                        # print(f"Checking placeholder: {placeholder} with value: {replacement_value}")

                        if placeholder in cell.text:
                            # Replace placeholder in the table cell text
                            cell.text = cell.text.replace(placeholder, replacement_value)

        # Save the modified document
        output_file = f"{output_folder}/Assignment_Schedule_{row['ID']}.docx"
        doc.save(output_file)
        print(f"Document saved: {output_file}")

    print(f"All assignment schedules have been saved to {output_folder}")


# Run the function
generate_assignment_schedules()
