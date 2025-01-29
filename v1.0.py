""""

# THIS VERSION 1 HAS BEEN COMMENTED OUT TO ALLOW FOR IMPROVEMENTS, SCROLL DOWN FOR VERSION 2

import pandas as pd
from docx import Document
import os

def generate_assignment_schedules(excel_file, template_file, output_folder):
    # Step 1: Load the Excel file
    df = pd.read_excel(excel_file)
    print(f"DataFrame preview:\n{df.head()}")  # Debug: show the first few rows

    # Step 2: Ensure the output folder exists
    os.makedirs(output_folder, exist_ok=True)

    # Step 3: Loop through each row in the Excel file
    for index, row in df.iterrows():
        print(f"Processing row {index}")  # Debug: show which row we are processing

        # Step 4: Load the Word template
        doc = Document(template_file)

        # Step 5: Replace placeholders in the document
        for column_name in df.columns:
            placeholder = f"{{{column_name}}}"  # Use double curly braces for the placeholder
            replacement_value = row[column_name]  # Access the value using index-based access
            print(f"Replacing placeholder '{placeholder}' with '{replacement_value}'")  # Debug log

            # Format date columns
            if isinstance(replacement_value, pd.Timestamp):
                replacement_value = replacement_value.strftime('%d/%m/%Y')

            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(replacement_value))

            # Replace placeholders in table cells
            for table in doc.tables:
                for table_row in table.rows:
                    for cell in table_row.cells:
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(replacement_value))

        # Step 6: Save the updated Word document
        output_file = os.path.join(
            output_folder, f"Assignment_Schedule_{row['{ID}']}.docx"
        )
        doc.save(output_file)
        print(f"Document saved: {output_file}")

    print(f"All assignment schedules have been saved to {output_folder}")

# Example usage
excel_file = r"C:\Users\petergrint.DPP_DOMAIN\PycharmProjects\ContractGenerator\input_data.xlsx"
template_file = r"C:\Users\petergrint.DPP_DOMAIN\PycharmProjects\ContractGenerator\template.docx"
output_folder = r"C:\Users\petergrint.DPP_DOMAIN\PycharmProjects\ContractGenerator\output_schedules"

generate_assignment_schedules(excel_file, template_file, output_folder)
"""